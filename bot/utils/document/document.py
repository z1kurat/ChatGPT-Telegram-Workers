import docx, os
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Document():
    def __init__(self):
        self.doc = docx.Document()
        section = self.doc.sections[0]
        section.left_margin = Mm(20.4)
        section.right_margin = Mm(10)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(10)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)

    def preprocessing(self, answer):
        # answer = answer.replace(' \n', '\n')
        return answer

    def filling(self, question, answer):
        para = self.doc.add_paragraph()
        bold_para = para.add_run(question)
        bold_para.bold = True
        para = self.doc.add_paragraph(self.preprocessing(answer=answer))
        # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print(f'Ответ на вопрос получен...')

    def saving(self, chat_id):
        self.doc.save(f'OUTPUT_DOCS\output_file{chat_id}.docx')
        # os.startfile('Answers_GPT-3.5-turbo.docx')